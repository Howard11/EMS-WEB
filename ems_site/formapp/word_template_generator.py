# -*- coding: utf-8 -*-
"""
Template-Based Word Generator (v6 - Comprehensive Field Mapping)

Fills data INTO '1.輔導單位基礎資料調查表v3.docx'.
Uses an exhaustive filling strategy for merged cells.
"""
from docx import Document
from docx.shared import Pt
import os

TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), '1.輔導單位基礎資料調查表v3.docx')

def _set_cell_text(cell, text, bold=False, size=10):
    """Set text in a cell safely."""
    try:
        if not cell.paragraphs:
            cell.add_paragraph()
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(str(text) if text is not None else '')
        run.font.size = Pt(size)
        run.bold = bold
    except Exception:
        pass

def _fill_range(rows, row_idx, start_col, end_col, text, bold=False, size=10):
    """Fill multiple cells in a range to ensure merged cells are covered."""
    for c in range(start_col, end_col + 1):
        try:
            _set_cell_text(rows[row_idx].cells[c], text, bold, size)
        except Exception:
            pass

def _safe_get(data, key, default=''):
    val = data.get(key, default)
    if val is None: return default
    if isinstance(val, list): return val[0] if val else default
    return str(val)

def _get_checkbox(condition):
    return "☑" if condition else "□"

def fill_table0(table, data):
    """Table 0: 機構基本資料 (Mapped to cells from inspection)"""
    rows = table.rows
    _fill_range(rows, 1, 1, 4, _safe_get(data, 'company_name')) # 公司名稱
    _fill_range(rows, 1, 6, 8, _safe_get(data, 'tax_id'))       # 統一編號
    _fill_range(rows, 2, 1, 8, _safe_get(data, 'company_address')) # 地址
    _fill_range(rows, 3, 1, 3, _safe_get(data, 'industry_type'))   # 產業類別
    _fill_range(rows, 3, 5, 8, _safe_get(data, 'building_type'))   # 建物分類
    
    cap = _safe_get(data, 'capital')
    _fill_range(rows, 4, 1, 2, f"{cap} 千元" if cap else "____ 千元")
    rev = _safe_get(data, 'revenue')
    _fill_range(rows, 4, 4, 6, f"{rev} 千元" if rev else "____ 千元")
    emp = _safe_get(data, 'employees')
    _fill_range(rows, 4, 8, 8, f"{emp} 人" if emp else "____ 人")
    
    _fill_range(rows, 5, 1, 2, _safe_get(data, 'contact_name'))
    _fill_range(rows, 5, 4, 5, _safe_get(data, 'contact_title'))
    _fill_range(rows, 5, 7, 8, _safe_get(data, 'phone'))
    _fill_range(rows, 6, 2, 8, _safe_get(data, 'email'))
    
    u = _safe_get(data, 'electricity_usage') or _safe_get(data, 'total_power_usage')
    c = _safe_get(data, 'electricity_cost') or _safe_get(data, 'total_power_cost')
    _fill_range(rows, 7, 1, 3, f"{u or '____'} 度/年")
    _fill_range(rows, 7, 5, 8, f"{c or '____'} 萬元/年")

    # Row 8: Non-electric usage
    gas = _safe_get(data, 'gas_usage'); lpg = _safe_get(data, 'lpg_usage')
    gaso = _safe_get(data, 'gasoline_usage'); die = _safe_get(data, 'diesel_usage')
    val_text = f"天然氣 {gas or '___'}m³ | LPG {lpg or '___'}噸 | 汽油 {gaso or '___'}公秉 | 柴油 {die or '___'}公秉"
    _fill_range(rows, 8, 1, 3, val_text)
    
    gc = _safe_get(data, 'gas_cost'); lc = _safe_get(data, 'lpg_cost')
    oc = _safe_get(data, 'gasoline_cost'); dc = _safe_get(data, 'diesel_cost')
    cost_text = f"天然氣 {gc or '___'}萬 | LPG {lc or '___'}萬 | 汽油 {oc or '___'}萬 | 柴油 {dc or '___'}萬"
    _fill_range(rows, 8, 5, 8, cost_text)

def fill_table1(table, data):
    """Table 1: 欲量測建物資料 (High complexity mapping)"""
    rows = table.rows
    # Row 1: Operating hours
    h = _safe_get(data, 'hours_per_day'); d = _safe_get(data, 'days_per_year')
    _fill_range(rows, 1, 0, 24, f"營運時數：{h or '24'} 小時/天   {d or '365'} 天/年")
    
    # Row 2: Building metadata
    _fill_range(rows, 2, 3, 10, _safe_get(data, 'building_name'))
    _fill_range(rows, 2, 11, 24, _safe_get(data, 'building_location'))
    
    # Row 3: Structure and Floors
    rc = _get_checkbox(_safe_get(data, 'building_form_rc') == 'on') + "RC"
    sc = _get_checkbox(_safe_get(data, 'building_form_sc') == 'on') + "SC"
    src = _get_checkbox(_safe_get(data, 'building_form_src') == 'on') + "SRC"
    _fill_range(rows, 3, 3, 8, f"{rc} {sc} {src}")
    _fill_range(rows, 3, 9, 14, _safe_get(data, 'floor_category'))
    _fill_range(rows, 3, 15, 20, _safe_get(data, 'floors'))
    _fill_range(rows, 3, 21, 24, _safe_get(data, 'beds') or _safe_get(data, 'partitions'))
    
    # Row 4: Areas
    _fill_range(rows, 4, 3, 10, f"{_safe_get(data, 'floor_area')} m²")
    _fill_range(rows, 4, 11, 20, f"{_safe_get(data, 'air_conditioning_area') or _safe_get(data, 'air_area')} m²")
    _fill_range(rows, 4, 21, 24, f"{_safe_get(data, 'parking_area')} m²")
    
    # Row 5: Contract
    _fill_range(rows, 5, 3, 10, f"{_safe_get(data, 'contract_capacity')} kW")
    _fill_range(rows, 5, 11, 24, f"{_safe_get(data, 'offpeak_contract_capacity')} kW")
    
    # Row 6: Demand
    _fill_range(rows, 6, 3, 10, f"{_safe_get(data, 'max_demand')} kW")
    _fill_range(rows, 6, 11, 24, f"{_safe_get(data, 'min_demand')} kW")
    
    # Row 7: Power Number & Type
    _fill_range(rows, 7, 3, 10, _safe_get(data, 'power_number'))
    types = [
        (_safe_get(data, 'auxiliary_table_lamp') == 'on', "表燈用電"),
        (_safe_get(data, 'auxiliary_peak_demand') == 'on', "低壓需量"),
        (_safe_get(data, 'auxiliary_high_voltage') == 'on', "高壓需量"),
        (_safe_get(data, 'auxiliary_extra_high_voltage') == 'on', "特高壓需量")
    ]
    type_str = " ".join([_get_checkbox(cond) + name for cond, name in types])
    _fill_range(rows, 7, 11, 24, type_str)
    
    # Row 8: Time Price
    p1 = _get_checkbox(_safe_get(data, 'time_pricing_one_stage') == 'on') + "二段式時間電價"
    p2 = _get_checkbox(_safe_get(data, 'time_pricing_two_stage') == 'on') + "三段式時間電價"
    _fill_range(rows, 8, 3, 24, f"{p1}   {p2}")
    
    # Row 9: Auto PF
    ps = _safe_get(data, 'auto_power_factor')
    val1 = _safe_get(data, 'power_factor_percentage')
    val2 = _safe_get(data, 'power_factor_percentage_no')
    pf_str = f"{_get_checkbox(ps == '有')}有，功率因數：{val1 or '____'} %    {_get_checkbox(ps == '無')}無，功率因數：{val2 or '____'} %"
    _fill_range(rows, 9, 3, 24, pf_str)
    
    # Row 10: Demand Management
    dm = _safe_get(data, 'demand_management')
    oth = _safe_get(data, 'demand_management_other')
    dm_text = f"{_get_checkbox(dm=='人工')}人工 {_get_checkbox(dm=='自動')}自動 {_get_checkbox(dm=='無')}無 {_get_checkbox(dm=='其他')}其他 {oth}"
    _fill_range(rows, 10, 3, 24, dm_text)
    
    # Row 11: Yearly Power
    u = _safe_get(data, 'total_power_usage') or _safe_get(data, 'electricity_usage')
    c = _safe_get(data, 'total_power_cost') or _safe_get(data, 'electricity_cost')
    _fill_range(rows, 11, 3, 10, f"{u or '___________'} 度/年")
    _fill_range(rows, 11, 11, 24, f"全年電力 使用費用: {c or '___________'} 萬元/年") # Special handle for merged label

    # Row 12: Non-electric detailed
    gas = _safe_get(data, 'gas_usage'); lpg = _safe_get(data, 'lpg_usage'); gaso = _safe_get(data, 'gasoline_usage'); die = _safe_get(data, 'diesel_usage')
    v_t = f"天然氣 {gas or '___'}m³ | LPG {lpg or '___'}噸 | 汽油 {gaso or '___'}公秉 | 柴油 {die or '___'}公秉"
    _fill_range(rows, 12, 3, 10, v_t)
    gc = _safe_get(data, 'gas_cost'); lc = _safe_get(data, 'lpg_cost'); oc = _safe_get(data, 'gasoline_cost'); dc = _safe_get(data, 'diesel_cost')
    c_t = f"天然氣 {gc or '___'}萬 | LPG {lc or '___'}萬 | 汽油 {oc or '___'}萬 | 柴油 {dc or '___'}萬"
    _fill_range(rows, 12, 17, 24, c_t)
    
    # Row 15: Power Distribution percentages
    cols = {'air':0, 'light':1, 'freezer':4, 'socket':7, 'fan':10, 'water':14, 'elevator':19, 'other':23}
    for k, idx in cols.items():
        _set_cell_text(rows[15].cells[idx], _safe_get(data, f'electric_{k}'))
        
    # Row 18-19: Heat Distribution
    h_map = {'boiler':2, 'generator':5, 'disinfection':8, 'kitchen':13, 'other1':18, 'other2':22}
    for k, idx in h_map.items():
        _set_cell_text(rows[18].cells[idx], _safe_get(data, f'fuel_{k}'))
        _set_cell_text(rows[19].cells[idx], _safe_get(data, f'heat_{k}'))

def fill_table2(table, data):
    """Table 2: Diagrams and ISO questions."""
    rows = table.rows
    diagrams = data.get('diagram[]', [])
    if isinstance(diagrams, str): diagrams = [diagrams]
    d_str = "(一)提供組織圖、配置圖 | " + " ".join([_get_checkbox(name in diagrams) + name for name in ["組織圖", "院區配置圖", "各樓層配置圖"]])
    _fill_range(rows, 0, 0, 0, d_str)
    
    iso = _safe_get(data, 'iso14064')
    iso_text = f"(二) 是否有進行ISO 14064組織型溫室氣體盤查? {_get_checkbox(iso=='有')}有  {_get_checkbox(iso=='無')}無"
    _fill_range(rows, 1, 0, 0, iso_text)
    
    an = _safe_get(data, 'area_name'); ey = _safe_get(data, 'emission_year'); ev = _safe_get(data, 'emission_value')
    em_text = f"若已有盤查，請填：全院區 {an or '____'} {ey or '____'}年排放 {ev or '____'} 公噸CO2e"
    _fill_range(rows, 2, 0, 0, em_text)
    _fill_range(rows, 4, 0, 0, "主要耗能設施清單 (詳細數據見文末表格)", bold=True)

def generate_from_template(json_data, output_path, template_path=None):
    if template_path is None: template_path = TEMPLATE_PATH
    doc = Document(template_path)
    
    if len(doc.tables) >= 1: fill_table0(doc.tables[0], json_data)
    if len(doc.tables) >= 2: fill_table1(doc.tables[1], json_data)
    if len(doc.tables) >= 3: fill_table2(doc.tables[2], json_data)
    
    # Append technical report as Section III
    from .word_generator import (
        add_lighting_system_table, add_air_conditioning_table, 
        add_ice_storage_table, add_boiler_system_tables, 
        add_heat_pump_table, add_elevator_system_table, add_air_compressor_table
    )
    
    doc.add_page_break()
    h = doc.add_heading('三、耗能系統詳細規格數據表 (網頁登錄完整內容)', level=1)
    h.alignment = 1
    
    add_lighting_system_table(doc, json_data)
    add_air_conditioning_table(doc, json_data)
    add_ice_storage_table(doc, json_data)
    add_boiler_system_tables(doc, json_data)
    add_heat_pump_table(doc, json_data)
    add_elevator_system_table(doc, json_data)
    add_air_compressor_table(doc, json_data)

    doc.save(output_path)
    return output_path
