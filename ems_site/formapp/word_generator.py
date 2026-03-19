# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json


def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders %s/>' % nsdecls('w'))
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = parse_xml('<%s %s/>' % (tag, nsdecls('w')))
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data: element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_cell_background(cell, fill):
    """Set cell background color"""
    from docx.oxml import OxmlElement
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def safe_add_table(doc, rows, cols):
    """Safely add a table with grid style."""
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = 'Table Grid'
    except Exception:
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={"sz": 4, "val": "single"}, bottom={"sz": 4, "val": "single"},
                                start={"sz": 4, "val": "single"}, end={"sz": 4, "val": "single"})
    return table

def _safe_get(data, key, default=''):
    """Safe key retrieval."""
    val = data.get(key, default)
    if val is None: return default
    if isinstance(val, list): return val[0] if val and val[0] is not None else default
    return str(val)

def _get_checkbox(condition):
    return "☑" if condition else "□"

def _get_data_list(data, key):
    """Get a list of values for a key, handling single values too."""
    v = data.get(key, [])
    if isinstance(v, list): return v
    return [v] if v is not None else []

def _fill_range(rows, row_idx, start_col, end_col, text, bold=False, size=10, background=None, align=None):
    """Helper to fill text into a range of cells (merged or discrete)."""
    cell = rows[row_idx].cells[start_col]
    if start_col < end_col: cell.merge(rows[row_idx].cells[end_col])
    if not cell.paragraphs: cell.add_paragraph()
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(str(text) if text is not None else '')
    run.font.size = Pt(size); run.bold = bold
    if align: p.alignment = align
    if background: set_cell_background(cell, background)
    return cell

def _fill_block(table, r1, c1, r2, c2, text, bold=False, size=10, background=None, align=None):
    """Fill a rectangular block of cells."""
    cell = table.cell(r1, c1); cell.merge(table.cell(r2, c2))
    if not cell.paragraphs: cell.add_paragraph()
    p = cell.paragraphs[0]; p.clear()
    run = p.add_run(str(text) if text is not None else '')
    run.font.size = Pt(size); run.bold = bold
    if align: p.alignment = align
    if background: set_cell_background(cell, background)
    return cell


def add_basic_info_section(doc, data):
    """Section 1: 訪視機構基本資料"""
    doc.add_heading('1、訪視機構基本資料表', level=2)
    table = safe_add_table(doc, rows=18, cols=9); rows = table.rows
    _fill_range(rows, 0, 0, 8, '基本資料表', bold=True, background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(rows, 1, 0, 0, '公司名稱', background='D9E2F3'); _fill_range(rows, 1, 1, 4, _safe_get(data, 'company_name'))
    _fill_range(rows, 1, 5, 5, '統一編號', background='D9E2F3'); _fill_range(rows, 1, 6, 8, _safe_get(data, 'tax_id'))
    _fill_range(rows, 2, 0, 0, '公司地址', background='D9E2F3'); _fill_range(rows, 2, 1, 8, _safe_get(data, 'company_address'))
    _fill_range(rows, 3, 0, 0, '產業類別', background='D9E2F3'); _fill_range(rows, 3, 1, 4, _safe_get(data, 'industry_type'))
    _fill_range(rows, 3, 5, 5, '建物分類', background='D9E2F3'); _fill_range(rows, 3, 6, 8, _safe_get(data, 'building_type'))
    _fill_range(rows, 4, 0, 0, '資本額', background='D9E2F3'); _fill_range(rows, 4, 1, 1, f"{_safe_get(data, 'capital')} 千元")
    _fill_range(rows, 4, 2, 3, '營業額', background='D9E2F3'); _fill_range(rows, 4, 4, 5, f"{_safe_get(data, 'revenue')} 千元")
    _fill_range(rows, 4, 6, 6, '員工人數', background='D9E2F3'); _fill_range(rows, 4, 7, 8, f"{_safe_get(data, 'employees')} 人")
    _fill_range(rows, 5, 0, 0, '聯絡人', background='D9E2F3'); _fill_range(rows, 5, 1, 2, _safe_get(data, 'contact_name'))
    _fill_range(rows, 5, 3, 3, '職稱', background='D9E2F3'); _fill_range(rows, 5, 4, 5, _safe_get(data, 'contact_title'))
    _fill_range(rows, 5, 6, 6, '電話', background='D9E2F3'); _fill_range(rows, 5, 7, 8, _safe_get(data, 'phone'))
    _fill_range(rows, 6, 0, 1, 'E-mail', background='D9E2F3'); _fill_range(rows, 6, 2, 8, _safe_get(data, 'email'))
    _fill_range(rows, 7, 0, 1, '全區電力\n使用總量', background='D9E2F3'); _fill_range(rows, 7, 2, 4, f"{_safe_get(data, 'total_power_usage')} 度/年")
    _fill_block(table, 7, 5, 7, 6, '全區電力\n使用費用', background='D9E2F3'); _fill_range(rows, 7, 7, 8, f"{_safe_get(data, 'total_power_cost')} 萬元/年")
    _fill_block(table, 8, 0, 13, 1, '非電力能源\n使用總量', background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 8, 5, 13, 6, '非電力能源\n使用費用', background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    fuels = [('天然氣', 'gas_usage', 'm³', 'gas_cost'), ('液化石油氣', 'lpg_usage', '噸', 'lpg_cost'), ('汽油', 'gasoline_usage', '公秉', 'gasoline_cost'), ('柴油', 'diesel_usage', '公秉', 'diesel_cost'), ('其他1', 'other_usage_1', '', 'other_cost_1'), ('其他2', 'other_usage_2', '', 'other_cost_2')]
    for i, (name, uk, unit, ck) in enumerate(fuels):
        r = 8+i; _fill_range(rows, r, 2, 2, name); _fill_range(rows, r, 3, 4, f"{_safe_get(data, uk)} {unit}"); _fill_range(rows, r, 7, 7, name); _fill_range(rows, r, 8, 8, f"{_safe_get(data, ck)} 萬元/年")
    _fill_block(table, 14, 0, 17, 1, '各類能源熱值', background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    h_labels = [('天然氣', 'gas_heat'), ('液化石油氣', 'lpg_heat'), ('汽油', 'gasoline_heat'), ('柴油', 'diesel_heat')]
    for i, (name, hk) in enumerate(h_labels):
        r = 14+i; _fill_range(rows, r, 2, 2, name); _fill_range(rows, r, 3, 8, _safe_get(data, hk))
    doc.add_paragraph()


def add_building_comprehensive_section(doc, data):
    """Section 2: 欲量測建物資料 (ISO 14064 + Metadata)."""
    doc.add_heading('2、欲量測建物各種能系統基本資料', level=2)
    t1 = safe_add_table(doc, rows=1, cols=1); diag = data.get('diagram[]', []) if isinstance(data.get('diagram[]'), list) else [data.get('diagram[]', '')]
    _fill_range(t1.rows, 0, 0, 0, f"(一)請貴院協助提供院所組織圖、及欲測建物各樓層配置圖\n{_get_checkbox('組織圖' in diag)}組織圖  {_get_checkbox('院區配置圖' in diag)}院區配置圖  {_get_checkbox('各樓層配置圖' in diag)}各樓層配置圖")
    doc.add_paragraph()
    t2 = safe_add_table(doc, rows=2, cols=1); iso = _safe_get(data, 'iso14064')
    _fill_range(t2.rows, 0, 0, 0, f"(二) 是否有進行ISO 14064組織型溫室氣體盤查？ {_get_checkbox(iso=='無')}無  {_get_checkbox(iso=='有')}有")
    _fill_range(t2.rows, 1, 0, 0, f"若已有14064組織型溫室氣體盤查，請填\n全院區 {(_safe_get(data, 'area_name')) or '____'} 年排放 {_safe_get(data, 'emission_year') or '____'} 公噸CO2e")
    doc.add_paragraph()
    _fill_range(doc.add_table(rows=1, cols=1).rows, 0, 0, 0, "(三)欲量測建物或場域能源使用基礎資料", bold=True)
    table = safe_add_table(doc, rows=25, cols=24); rows = table.rows
    _fill_range(rows, 0, 0, 23, '欲量測建物資料', bold=True, background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(rows, 1, 0, 23, f"營運時數： {_safe_get(data, 'hours_per_day') or '24'} 小時/天  {_safe_get(data, 'days_per_year') or '365'} 天/年")
    _fill_range(rows, 2, 0, 3, '建物名稱', background='D9E2F3'); _fill_range(rows, 2, 4, 11, _safe_get(data, 'building_name')); _fill_range(rows, 2, 12, 15, '建物場址', background='D9E2F3'); _fill_range(rows, 2, 16, 23, _safe_get(data, 'building_location'))
    _fill_range(rows, 3, 0, 3, '建物形式', background='D9E2F3')
    f_str = f"{_get_checkbox(_safe_get(data, 'building_form_rc')=='on')}RC {_get_checkbox(_safe_get(data, 'building_form_sc')=='on')}SC {_get_checkbox(_safe_get(data, 'building_form_src')=='on')}SRC"
    _fill_range(rows, 3, 4, 11, f_str); _fill_range(rows, 3, 12, 14, '樓層別', background='D9E2F3'); _fill_range(rows, 3, 15, 17, _safe_get(data, 'floor_category')); _fill_range(rows, 3, 18, 19, '樓層數', background='D9E2F3'); _fill_range(rows, 3, 20, 21, _safe_get(data, 'floors')); _fill_range(rows, 3, 22, 22, '隔間數', background='D9E2F3'); _fill_range(rows, 3, 23, 23, _safe_get(data, 'beds') or _safe_get(data, 'partitions'))
    _fill_range(rows, 4, 0, 3, '樓地板面積', background='D9E2F3'); _fill_range(rows, 4, 4, 7, f"{_safe_get(data, 'floor_area')} m²"); _fill_range(rows, 4, 8, 11, '空調使用面積', background='D9E2F3'); _fill_range(rows, 4, 12, 15, f"{_safe_get(data, 'air_conditioning_area') or _safe_get(data, 'air_area')} m²"); _fill_range(rows, 4, 16, 19, '停車面積', background='D9E2F3'); _fill_range(rows, 4, 20, 23, f"{_safe_get(data, 'parking_area')} m²")
    
    # Restored Rows 5-10
    _fill_range(rows, 5, 0, 3, '契約容量', background='D9E2F3'); _fill_range(rows, 5, 4, 11, f"{_safe_get(data, 'contract_capacity')} kW"); _fill_range(rows, 5, 12, 15, '離峰契約容量', background='D9E2F3'); _fill_range(rows, 5, 16, 23, f"{_safe_get(data, 'offpeak_contract_capacity')} kW")
    _fill_range(rows, 6, 0, 3, '最高需量', background='D9E2F3'); _fill_range(rows, 6, 4, 11, f"{_safe_get(data, 'max_demand')} kW"); _fill_range(rows, 6, 12, 15, '最低需量', background='D9E2F3'); _fill_range(rows, 6, 16, 23, f"{_safe_get(data, 'min_demand')} kW")
    _fill_range(rows, 7, 0, 3, '台電電號', background='D9E2F3'); _fill_range(rows, 7, 4, 11, _safe_get(data, 'power_number')); _fill_range(rows, 7, 12, 15, '用電類型', background='D9E2F3')
    types = [(_safe_get(data, 'auxiliary_table_lamp')=='on', '表燈用電'), (_safe_get(data, 'auxiliary_peak_demand')=='on', '低壓需量'), (_safe_get(data, 'auxiliary_high_voltage')=='on', '高壓需量'), (_safe_get(data, 'auxiliary_extra_high_voltage')=='on', '特高壓')]
    _fill_range(rows, 7, 16, 23, " ".join([_get_checkbox(c)+n for c,n in types]))
    _fill_range(rows, 8, 0, 3, '時間電價', background='D9E2F3'); _fill_range(rows, 8, 4, 23, f"{_get_checkbox(_safe_get(data, 'time_pricing_one_stage')=='on')}二段式時間電價  {_get_checkbox(_safe_get(data, 'time_pricing_two_stage')=='on')}三段式時間電價")
    _fill_range(rows, 9, 0, 3, '自動功因', background='D9E2F3'); _fill_range(rows, 9, 4, 23, f"{_get_checkbox(_safe_get(data, 'auto_power_factor')=='有')}有，功率因數：{_safe_get(data, 'power_factor_percentage') or '____'}%    {_get_checkbox(_safe_get(data, 'auto_power_factor')=='無')}無，功率因數：{_safe_get(data, 'power_factor_percentage_no') or '____'}%")
    _fill_range(rows, 10, 0, 3, '需量管理', background='D9E2F3'); _fill_range(rows, 10, 4, 23, " ".join([_get_checkbox(_safe_get(data, 'demand_management')==l)+l for l in ['人工', '自動', '無', '其他']]) + f" {_safe_get(data, 'demand_management_other')}")

    _fill_range(rows, 11, 0, 3, '全年電力\n使用總量', background='D9E2F3'); _fill_range(rows, 11, 4, 11, f"{_safe_get(data, 'total_power_usage')} 度/年"); _fill_range(rows, 11, 12, 15, '全年電力\n使用費用', background='D9E2F3'); _fill_range(rows, 11, 16, 23, f"{_safe_get(data, 'total_power_cost')} 萬元/年")
    _fill_block(table, 12, 0, 17, 3, '非電力能源\n使用總量', background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 12, 12, 17, 15, '非電力能源\n使用費用', background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    fuels = [('天然氣', 'gas_usage', 'm³', 'gas_cost'), ('液化石油氣', 'lpg_usage', '噸', 'lpg_cost'), ('汽油', 'gasoline_usage', '公秉', 'gasoline_cost'), ('柴油', 'diesel_usage', '公秉', 'diesel_cost'), ('其他1', 'other_usage_1', '', 'other_cost_1'), ('其他2', 'other_usage_2', '', 'other_cost_2')]
    for i, (name, uk, unit, ck) in enumerate(fuels):
        r = 12+i; _fill_range(rows, r, 4, 6, name); _fill_range(rows, r, 7, 11, f"{_safe_get(data, uk)} {unit}"); _fill_range(rows, r, 16, 18, name); _fill_range(rows, r, 19, 23, f"{_safe_get(data, ck)} 萬元/年")
    _fill_block(table, 18, 0, 18, 23, '電力使用分佈占比(%)*', bold=True, background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    k_e = ['electric_air', 'electric_light', 'electric_freezer', 'electric_socket', 'electric_fan', 'electric_hot_water', 'electric_elevator', 'electric_other']
    h_e = ['空調', '照明', '冷凍冷藏', '插座', '送排風', '給排水', '電(扶)梯', '其他']
    for i, h in enumerate(h_e): _fill_range(rows, 19, i*3, i*3+2, h, background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER); _fill_range(rows, 20, i*3, i*3+2, _safe_get(data, k_e[i]), align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 21, 0, 21, 23, '熱能使用分佈占比(%)*', bold=True, background='D9E2F3', align=WD_ALIGN_PARAGRAPH.CENTER)
    w_h = [3, 4, 3, 3, 3, 4, 4]; h_h = ['設備', '鍋爐設備', '發電設備', '消毒設備', '廚房', '其他1', '其他2']; h_f = ['fuel_boiler', 'fuel_generator', 'fuel_disinfection', 'fuel_kitchen', 'fuel_other1', 'fuel_other2']; h_p = ['heat_boiler', 'heat_generator', 'heat_disinfection', 'heat_kitchen', 'heat_other1', 'heat_other2']
    curr = 0; 
    for i, w in enumerate(w_h):
        s = curr; e = curr + w - 1; curr += w; _fill_range(rows, 22, s, e, h_h[i], background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
        if i == 0: _fill_range(rows, 23, s, e, '使用燃料', background='F2F2F2'); _fill_range(rows, 24, s, e, '占比(%)', background='F2F2F2')
        else: _fill_range(rows, 23, s, e, _safe_get(data, h_f[i-1]), align=WD_ALIGN_PARAGRAPH.CENTER); _fill_range(rows, 24, s, e, _safe_get(data, h_p[i-1]), align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph('註：電力使用分佈占比及熱能使用分佈占比, 若無實際數值, 請以營運狀態進行預估填寫')


def add_lighting_system_v5(doc, data):
    """Section 3: 照明系統."""
    doc.add_paragraph('■主要耗能設施').bold = True; doc.add_heading('1.照明系統', level=3)
    table = safe_add_table(doc, rows=14, cols=8)
    _fill_block(table, 0, 0, 1, 0, '燈具種類\n(日光燈、\n省電燈泡、\nLED…)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(table.rows, 0, 1, 3, '燈具規格', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 0, 4, 1, 4, '燈具\n電力功率值\n(瓦/具)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 0, 5, 1, 5, '數量', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 0, 6, 0, 6, '運轉', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER); _fill_block(table, 1, 6, 1, 6, '時數\n(時/日)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_block(table, 0, 7, 0, 7, '運作', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER); _fill_block(table, 1, 7, 1, 7, '天數\n(日/年)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(table.rows, 1, 1, 1, '燈管型式\n(T5、T8、螺\n旋、平板…)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(table.rows, 1, 2, 2, '容量規格(\nW)', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    _fill_range(table.rows, 1, 3, 3, '安定器型式', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER)
    light_keys = ['light_type[]', 'light_spec_type[]', 'light_spec_cap[]', 'light_spec_ballast[]', 'light_power[]', 'light_count[]', 'light_hours[]', 'light_days[]']
    col_data = [_get_data_list(data, k) for k in light_keys]
    num_rows = min(max(len(d) for d in col_data) if col_data else 0, 12)
    for i in range(num_rows):
        for j in range(8):
            val = col_data[j][i] if i < len(col_data[j]) else ''; table.cell(i + 2, j).text = str(val if val is not None else '')
    doc.add_paragraph(); rule = _safe_get(data, 'light_saving'); desc = _safe_get(data, 'light_saving_desc'); tf = safe_add_table(doc, rows=1, cols=2)
    _fill_range(tf.rows, 0, 0, 0, '照明節能規範', background='F2F2F2', align=WD_ALIGN_PARAGRAPH.CENTER); _fill_range(tf.rows, 0, 1, 1, f"{_get_checkbox(rule=='無')}無    {_get_checkbox(rule=='有')}有, 說明：{desc}")
    doc.add_paragraph()


def add_air_conditioning_v4(doc, data):
    """Section 4: 空調系統."""
    doc.add_heading('2.空調系統(冰水主機)', level=3); no_ac = '無空調系統' in data.get('diagram[]', [])
    _fill_range(doc.add_table(rows=1, cols=1).rows, 0, 0, 0, f"{_get_checkbox(no_ac)}無空調系統")
    items = [('① 空調主機編號', 'ac_id[]', 't'), ('空調/冰水機型式', 'ac_type[]', 'type'), ('壓縮機類型', 'ac_comp_type[]', 'comp'), ('② 設備數量', 'ac_count[]', 't'), ('③ 額定能力(RT)', 'ac_capacity[]', 't'), ('④ 功率(kW)', 'ac_power_v[]', 't'), ('⑤ 設備廠牌', 'ac_brand[]', 't'), ('⑥ 設備型號', 'ac_model[]', 't'), ('⑦ 採購年份', 'ac_buy_year[]', 't'), ('⑧ 冷卻能力', 'ac_cool_cap[]', 't'), ('⑨ 壓縮機數量', 'ac_comp_count[]', 't'), ('運轉時數(h/y)', 'ac_hours[]', 't'), ('冷媒種類', 'ac_refrigerant[]', 't'), ('冰水出水溫度', 'ac_out_temp[]', 't'), ('負載率(%)', 'ac_load[]', 't'), ('供應樓層總面積', 'ac_floor_area[]', 't'), ('水塔風扇變頻控制', 'ac_vfd[]', 't')]
    table = safe_add_table(doc, rows=len(items), cols=6)
    for r_idx, (lab, key, ui) in enumerate(items):
        _fill_range(table.rows, r_idx, 0, 0, lab, background='F2F2F2'); vals = _get_data_list(data, key)
        for col in range(1, 6):
            v = vals[col-1] if col-1 < len(vals) else ''
            if ui == 'type': table.cell(r_idx, col).text = f"{_get_checkbox(v=='氣冷式')}氣冷\n{_get_checkbox(v=='水冷式')}水冷"
            elif ui == 'comp': table.cell(r_idx, col).text = f"{_get_checkbox(v=='螺旋式')}螺旋\n{_get_checkbox(v=='離心式')}離心\n{_get_checkbox(v=='渦卷式')}渦卷"
            else: table.cell(r_idx, col).text = str(v if v is not None else '')
    doc.add_paragraph()


def add_ice_storage_v4(doc, data):
    """Section 4.1: 儲冰系統."""
    doc.add_heading('2.1儲冰系統', level=3); items = [('① 儲冰主機編號', 'ice_id[]', 't'), ('形式', 'ice_type[]', 'comp'), ('② 設備數量', 'ice_count[]', 't'), ('③ 額定能力(RT)', 'ice_capacity[]', 't'), ('④ 設備廠牌', 'ice_brand[]', 't'), ('⑤ 設備型號', 'ice_model[]', 't'), ('⑥ 採購年份', 'ice_buy_year[]', 't'), ('負載率(%)', 'ice_load[]', 't'), ('冰水溫度(°C)', 'ice_temp[]', 't'), ('運轉時數(h/y)', 'ice_hours[]', 't'), ('運作天數(d/y)', 'ice_days[]', 't')]
    table = safe_add_table(doc, rows=len(items), cols=6)
    for r_idx, (lab, key, ui) in enumerate(items):
        _fill_range(table.rows, r_idx, 0, 0, lab, background='F2F2F2'); vals = _get_data_list(data, key)
        for col in range(1, 6):
            v = vals[col-1] if col-1 < len(vals) else ''
            if ui == 'comp': table.cell(r_idx, col).text = f"{_get_checkbox(v=='螺旋式')}螺旋\n{_get_checkbox(v=='離心式')}離心\n{_get_checkbox(v=='渦卷式')}渦卷"
            else: table.cell(r_idx, col).text = str(v)
    doc.add_paragraph()


def add_boiler_system_v4(doc, data):
    """Section 5: 鍋爐系統."""
    doc.add_heading('3.鍋爐系統', level=3); doc.add_paragraph(f"■系統運轉時間：{_safe_get(data, 'boiler_running_hours')} 小時/日   運作天數：{_safe_get(data, 'boiler_running_days')} 日/年")
    table = safe_add_table(doc, rows=12, cols=4); _fill_range(table.rows, 0, 0, 0, '①鍋爐型式', background='F2F2F2')
    for i in range(1, 4):
        v = _safe_get(data, f'boiler_type_{i}')
        table.cell(0, i).text = f"{_get_checkbox(v=='水管式')}水管\n{_get_checkbox(v=='煙管式')}煙管\n{_get_checkbox(v=='貫流式')}貫流"
    labels = [('②鍋爐規格(T/H)', 'boiler_capacity_'), ('③設備數量(台數)', 'boiler_count_'), ('④額定功率(kW)', 'boiler_power_'), ('⑤排氣溫度(℃)', 'boiler_exhaust_temp_'), ('⑥設置年份(年)', 'boiler_year_'), ('⑦水量(噸/年)', 'boiler_water_'), ('⑧入水溫度(℃)', 'boiler_in_temp_'), ('⑨出水溫度(℃)', 'boiler_out_temp_')]
    for r, (l, k) in enumerate(labels):
        _fill_range(table.rows, r+1, 0, 0, l, background='F2F2F2')
        for i in range(1, 4): table.cell(r+1, i).text = _safe_get(data, f"{k}{i}")
    doc.add_paragraph()


def add_heat_pump_v4(doc, data):
    """Section 6: 熱泵系統."""
    doc.add_heading('3.1熱泵系統', level=3); table = safe_add_table(doc, rows=8, cols=4)
    items = [('年熱水需求(噸/年)', 'hp_annual_demand[]'), ('入/出水溫(°C)', 'hp_in_temp[]'), ('設置年份(年)', 'hp_year[]'), ('平均燃料單價', 'hp_fuel_price[]'), ('額定功率(kW)', 'hp_rated_power[]'), ('製熱能力(BTU/H)', 'hp_heating_capacity[]'), ('運轉時數(h/d)', 'hp_hours[]'), ('運作天數(d/y)', 'hp_days[]')]
    for r, (l, k) in enumerate(items):
        _fill_range(table.rows, r, 0, 0, l, background='F2F2F2'); vals = _get_data_list(data, k)
        for i in range(1, 4): table.cell(r, i).text = vals[i-1] if i-1 < len(vals) else ''
    doc.add_paragraph()


def add_elevator_v4(doc, data):
    """Section 7: 電梯系統."""
    doc.add_heading('4.電梯系統', level=3); table = safe_add_table(doc, rows=14, cols=6)
    labels = [('使用率(%)', 'elevator_usage_'), ('載重(kg)', 'elevator_weight_'), ('馬達效率(IE2/3/4)', 'elevator_efficiency_'), ('數量(台)', 'elevator_count_'), ('動力電源', 'elevator_power_'), ('主機功率(kW)', 'elevator_main_power_'), ('馬力(HP)', 'elevator_hp_'), ('能效(cfm/HP)', 'elevator_main_eff_'), ('速度(m/min)', 'elevator_speed_'), ('變壓器(KVA)', 'elevator_transformer_'), ('遮斷器(A)', 'elevator_controller_'), ('線徑(mm2)', 'elevator_cable_'), ('運轉時數(h/d)', 'elevator_hours_'), ('運作天數(d/y)', 'elevator_days_')]
    for r, (l, k) in enumerate(labels):
        if r >= 14: break
        _fill_range(table.rows, r, 0, 0, l, background='F2F2F2')
        for i in range(1, 6): table.cell(r, i).text = _safe_get(data, f"{k}{i}[]")
    doc.add_paragraph()


def add_compressor_v4(doc, data):
    """Section 8: 空壓系統."""
    doc.add_heading('5.空壓系統', level=3); table = safe_add_table(doc, rows=8, cols=6)
    labels = [('使用率(%)', 'compressor_usage_'), ('馬力(hp)', 'compressor_hp_'), ('馬達效率', 'compressor_efficiency_'), ('台數', 'compressor_count_'), ('平均負載(%)', 'compressor_load_'), ('能效(cfm/hp)', 'compressor_main_eff_'), ('運轉時數(h/d)', 'compressor_hours_'), ('運作天數(d/y)', 'compressor_days_')]
    for r, (l, k) in enumerate(labels):
        _fill_range(table.rows, r, 0, 0, l, background='F2F2F2')
        for i in range(1, 6): table.cell(r, i).text = _safe_get(data, f"{k}{i}[]")
    doc.add_paragraph()


def generate_energy_system_word(json_data, output_path):
    """Entry Point."""
    doc = Document()
    for s in doc.sections:
        s.page_width = Cm(21.0); s.page_height = Cm(29.7)
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5); s.left_margin = Cm(1.5); s.right_margin = Cm(1.5)
    title = doc.add_heading('能源管理系統現況調查報告', 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_basic_info_section(doc, json_data)
    add_building_comprehensive_section(doc, json_data)
    add_lighting_system_v5(doc, json_data)
    add_air_conditioning_v4(doc, json_data)
    add_ice_storage_v4(doc, json_data)
    add_boiler_system_v4(doc, json_data)
    add_heat_pump_v4(doc, json_data)
    add_elevator_v4(doc, json_data)
    add_compressor_v4(doc, json_data)
    doc.save(output_path); return output_path
