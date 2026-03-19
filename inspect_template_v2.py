import sys
import os
from docx import Document

def analyze_template():
    path = r'C:\Users\howar\OneDrive\Desktop\工作\EMS\Django\EMS_Web\EMS_Web\1.輔導單位基礎資料調查表v3.docx'
    if not os.path.exists(path):
        print(f"Error: File not found at {path}")
        return

    doc = Document(path)
    print(f"Total Tables: {len(doc.tables)}")
    
    for t_idx, table in enumerate(doc.tables):
        print(f"\n--- TABLE {t_idx} (Rows: {len(table.rows)}) ---")
        for r_idx, row in enumerate(table.rows):
            cell_texts = []
            for c_idx, cell in enumerate(row.cells):
                # Clean text for readability
                text = cell.text.replace('\n', ' ').strip()
                cell_texts.append(f"[{c_idx}]{text}")
            print(f"Row {r_idx}: {cell_texts}")

if __name__ == "__main__":
    analyze_template()
