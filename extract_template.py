# -*- coding: utf-8 -*-
"""Extract table structure from Word template and save to file."""
from docx import Document
import json

TEMPLATE_PATH = r'C:\Users\howar\OneDrive\Desktop\工作\EMS\Django\EMS_Web\EMS_Web\1.輔導單位基礎資料調查表v3.docx'
OUTPUT_PATH = r'C:\Users\howar\OneDrive\Desktop\工作\EMS\Django\EMS_Web\EMS_Web\template_structure.txt'

doc = Document(TEMPLATE_PATH)

with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
    f.write(f"Total tables: {len(doc.tables)}\n\n")
    
    for idx, table in enumerate(doc.tables):
        f.write(f"=== Table {idx}: {len(table.rows)} rows x {len(table.columns)} cols ===\n")
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for cell in row.cells:
                text = cell.text.strip().replace('\n', ' | ')[:50]
                cells_text.append(text)
            f.write(f"  Row {row_idx}: {cells_text}\n")
        f.write("\n")
        
print(f"Template structure saved to: {OUTPUT_PATH}")
